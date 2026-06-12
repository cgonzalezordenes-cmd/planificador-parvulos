import logging
import json
import os
import io
import re
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import CellIsRule

import azure.functions as func
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT
from groq import Groq

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

# ═══════════════════════════════════════════════════════
# PROMPTS
# ═══════════════════════════════════════════════════════

SYSTEM_PROMPT = """Eres una experta en Educación Parvularia chilena. Generas planificaciones semanales co-construidas y diversificadas para Nivel Medio Menor (3 años) siguiendo las Bases Curriculares de Educación Parvularia 2018.

Para cada dia elige el Ambito, Nucleo y OA mas adecuado segun el tema e ideas proporcionadas. Debes variar los ambitos entre los dias de la semana segun corresponda al tema. NUNCA uses el mismo Ambito mas de 2 veces en la semana.

A continuacion encontraras el listado OFICIAL y COMPLETO de Objetivos de Aprendizaje para Nivel Medio Menor (Segundo Nivel) del BCEP 2018. Usa EXACTAMENTE estos textos al seleccionar el OA y OAT:

═══════════════════════════════════════════
AMBITO 1: Desarrollo Personal y Social
═══════════════════════════════════════════

NUCLEO: Identidad y Autonomia
OA 1: Representar verbal y corporalmente diferentes emociones y sentimientos, en sus juegos.
OA 2: Manifestar disposicion y confianza al separarse de los adultos significativos.
OA 3: Reconocer en si mismo, en otras personas y en personajes de cuentos, emociones.
OA 4: Manifestar disposicion para regular sus emociones y sentimientos.
OA 5: Manifestar sus preferencias cuando participa o solicita participar en situaciones cotidianas y juegos.
OA 6: Actuar con progresiva independencia, ampliando su repertorio de acciones.
OA 7: Comunicar algunos rasgos de su identidad, como su nombre, caracteristicas corporales, genero.
OA 8: Apreciar sus caracteristicas identitarias, fortalezas y habilidades.
OA 9: Manifestar progresiva independencia en practicas de alimentacion, vestimenta, higiene corporal, bucal.
OA 10: Manifestar satisfaccion y confianza por su autovalia, comunicando algunos desafios alcanzados.
OA 11: Identificar alimentos que se consumen en algunas celebraciones propias de su familia y comunidad.
OA 12: Representar sus pensamientos y experiencias, atribuyendo significados a objetos o elementos.

NUCLEO: Convivencia y Ciudadania
OA 1: Participar en actividades y juegos grupales con sus pares, conversando, intercambiando pertenencias, cooperando.
OA 2: Disfrutar de instancias de interaccion social con diversas personas de la comunidad.
OA 3: Colaborar en situaciones cotidianas y de juego, proponiendo acciones simples frente a necesidades de pares.
OA 4: Colaborar en actividades, conmemoraciones o celebraciones culturales de su familia y comunidad.
OA 5: Iniciarse en la resolucion pacifica de conflictos, dialogando respecto de la situacion.
OA 6: Manifestar disposicion para practicar acuerdos de convivencia basica que regulan situaciones cotidianas.
OA 7: Identificar objetos, comportamientos y situaciones de riesgo que pueden atentar contra su seguridad.
OA 8: Reconocer acciones correctas e incorrectas para la convivencia armonica del grupo.
OA 9: Manifestar interes por algunos de sus derechos, tales como: ser escuchados, tener un nombre, jugar.
OA 10: Manifestar interes para interactuar con ninos y ninas, reconociendo la diversidad de sus caracteristicas.

NUCLEO: Corporalidad y Movimiento
OA 1: Reconocer situaciones en que se siente comodo corporalmente.
OA 2: Reconocer las principales partes, caracteristicas fisicas de su cuerpo y sus funciones.
OA 3: Experimentar diversas posibilidades de accion con su cuerpo.
OA 4: Reconocer el bienestar que le produce el movimiento libre en situaciones cotidianas y ludicas.
OA 5: Perfeccionar su coordinacion visomotriz fina, a traves del uso de diversos objetos, juguetes y utensilios.
OA 6: Adquirir control y equilibrio en movimientos, posturas y desplazamientos.
OA 7: Resolver desafios practicos en situaciones cotidianas y juegos, incorporando mayor precision y coordinacion.
OA 8: Utilizar categorias de ubicacion espacial y temporal (adelante/atras, arriba/abajo, adentro/afuera).

═══════════════════════════════════════════
AMBITO 2: Comunicacion Integral
═══════════════════════════════════════════

NUCLEO: Lenguaje Verbal
OA 1: Expresarse oralmente, empleando estructuras oracionales simples y respetando patrones gramaticales basicos, en distintas situaciones cotidianas y juegos.
OA 2: Comprender mensajes simples como instrucciones explicitas, explicaciones y preguntas relativas a objetos, personas, acciones, tiempo y lugar, identificando la intencionalidad comunicativa de diversos interlocutores.
OA 3: Identificar algunos atributos de los sonidos de diferentes fuentes sonoras como intensidad (fuerte/suave), velocidad (rapido/lento).
OA 4: Incorporar progresivamente nuevas palabras, al comunicar oralmente temas variados de su interes e informacion basica, en distintas situaciones cotidianas.
OA 5: Manifestar interes por descubrir el contenido de textos de diferentes formatos, a traves de la manipulacion, la exploracion, la escucha atenta y la formulacion de preguntas.
OA 6: Comprender a partir de la escucha atenta, contenidos explicitos de textos literarios y no literarios, reconociendo ideas centrales, senalando preferencias, realizando sencillas descripciones, preguntando sobre el contenido.
OA 7: Reconocer progresivamente el significado de diversas imagenes, logos, simbolos de su entorno cotidiano, en diversos soportes (incluye uso de TICs).
OA 8: Producir sus propios signos graficos en situaciones ludicas.

NUCLEO: Lenguajes Artisticos
OA 1: Manifestar interes por diversas producciones artisticas (arquitectura, modelado, piezas musicales, pintura, dibujos, titeres, obras de teatro, danzas, entre otras), describiendo algunas caracteristicas.
OA 2: Expresar sus preferencias, sensaciones y emociones relacionadas con diferentes recursos expresivos que se encuentran en sencillas obras visuales (colorido, formas), musicales (fuente, intensidad del sonido) o escenicas (desplazamiento, vestimenta, caracter expresivo).
OA 3: Interpretar canciones y juegos musicales, experimentando con diversos recursos tales como, la voz, el cuerpo, instrumentos musicales y objetos.
OA 4: Expresar corporalmente sensaciones y emociones experimentando con mimica, juegos teatrales, rondas, bailes y danzas.
OA 5: Expresar emociones, ideas y experiencias por medio de la plastica experimentando con recursos pictoricos, graficos y de modelado.
OA 6: Experimentar diversas posibilidades de expresion, combinando lenguajes artisticos en sus producciones.
OA 7: Representar a traves del dibujo, diversos elementos de su entorno, incorporando figuras cerradas, trazos intencionados y primeros esbozos de la figura humana.

═══════════════════════════════════════════
AMBITO 3: Interaccion y Comprension del Entorno
═══════════════════════════════════════════

NUCLEO: Exploracion del Entorno Natural
OA 1: Manifestar interes y asombro por diversos elementos, situaciones y fenomenos del entorno natural, explorando, observando, preguntando, describiendo, agrupando, entre otros.
OA 2: Comunicar verbalmente caracteristicas de elementos y paisajes de su entorno natural, tales como cuerpos celestes, cerros, desierto, flora; y de fenomenos como marejadas, sismos, tormentas, sequias.
OA 3: Descubrir que el sol es fuente de luz y calor para el planeta, a traves de experiencias directas o TICs.
OA 4: Comunicar algunas propiedades basicas de los elementos naturales que explora, tales como: colores, texturas, tamanios, temperaturas entre otras.
OA 5: Distinguir una variedad progresivamente mas amplia de animales y plantas, respecto a sus caracteristicas (tamanio, color, textura y morfologia), sus necesidades basicas y los lugares que habitan, al observarlos en forma directa, en libros ilustrados o en TICs.
OA 6: Colaborar en situaciones cotidianas, en acciones que contribuyen al desarrollo de ambientes sostenibles, tales como cerrar las llaves de agua, apagar aparatos electricos, entre otras.
OA 7: Emplear instrumentos y herramientas de observacion y recoleccion (lupas, frascos, recipientes, botellas, cucharas, embudos, pinzas, entre otros) en la exploracion del entorno natural.
OA 8: Experimentar mezclas y disoluciones con materiales cotidianos tales como: burbujas de jabon, agua salada, gelatina, describiendo los cambios observados.
OA 9: Reconocer que el aire y el agua son elementos vitales para las personas, los animales y las plantas, y que estos elementos pueden encontrarse con o sin contaminacion.

NUCLEO: Comprension del Entorno Sociocultural
OA 1: Describir actividades habituales de su comunidad, como ir de compras, jugar en la plaza, viajar en bus, entre otras, senalando su participacion en ellas.
OA 2: Describir caracteristicas de las formas de vida de su comunidad (viviendas, paisajes, costumbres), a traves de canciones, juegos, relatos y fotos familiares, entre otras.
OA 3: Seleccionar utensilios domesticos y objetos tecnologicos que les permiten resolver problemas en contextos sociales autenticos.
OA 4: Reconocer sucesos significativos de su historia personal y familiar, en diversas situaciones, tales como: conversaciones familiares, relatos de un agente comunitario, visitas a lugares, observacion de fotografias, entre otros.
OA 5: Identificar instituciones significativas de su entorno, describiendo actividades y rutinas representativas que en ellas se realizan.
OA 6: Identificar algunas normas de proteccion y seguridad de su entorno cotidiano referidas a alimentacion, transito y sismos, y otras pertinentes a su contexto geografico.
OA 7: Distinguir en paisajes de su localidad, elementos naturales (bosque, cerros, rios), y culturales (caminos, edificios, puentes).

NUCLEO: Pensamiento Matematico
OA 1: Reproducir patrones sonoros, visuales, gestuales, corporales u otros, de dos o tres elementos.
OA 2: Experimentar con diversos objetos, estableciendo relaciones al clasificar por dos atributos a la vez (forma, color, entre otros) y seriar por altura o longitud.
OA 3: Describir la posicion de objetos y personas, respecto de un punto u objeto de referencia, empleando conceptos de ubicacion y distancia tales como: dentro/fuera; encima/debajo; cerca/lejos.
OA 4: Orientarse temporalmente en situaciones cotidianas, mediante la utilizacion progresiva de algunas nociones y relaciones de secuencias, tales como: antes/despues, dia/noche, hoy/maniana.
OA 5: Emplear cuantificadores, tales como: mas/menos, mucho/poco, todo/ninguno, al comparar cantidades de objetos en situaciones cotidianas.
OA 6: Emplear progresivamente los numeros, para contar, identificar, cuantificar y comparar cantidades, hasta el 10 e indicar orden o posicion de algunos elementos en situaciones cotidianas o juegos.
OA 7: Representar progresivamente, numeros y cantidades en forma concreta y pictorica hasta el 10.
OA 8: Resolver progresivamente problemas simples, de manera concreta y pictorica, agregando o quitando hasta 5 elementos.
OA 9: Descubrir atributos de figuras 3D, mediante la exploracion de objetos presentes en su entorno.
OA 10: Identificar algunas acciones que se llevaron a cabo para resolver problemas.

El OAT siempre pertenece al Ambito "Desarrollo Personal y Social". Elige el nucleo y OAT mas complementario al OA principal elegido. IMPORTANTE: NO repitas el mismo ambito todos los dias — distribuye segun el tema de cada dia.

Responde SOLO con un objeto JSON valido sin texto adicional, sin bloques de codigo markdown.
El JSON debe tener claves: lunes, martes, miercoles, jueves, viernes. Cada una con exactamente estos campos:
{
  "titulo_dia": "Lunes 9 de junio: Titulo creativo de la experiencia",
  "ambito": "Nombre exacto del ambito del OA (Desarrollo Personal y Social | Comunicacion Integral | Interaccion y Comprension del Entorno)",
  "nucleo": "Nombre exacto del nucleo dentro del ambito (ej: Identidad y Autonomia, Lenguajes Artisticos, Pensamiento Matematico)",
  "oa_num": "Solo el numero del OA elegido (ej: OA 3)",
  "oa_texto": "Texto completo del OA elegido sin incluir el numero",
  "oat_nucleo": "Nombre del nucleo del OAT dentro de Desarrollo Personal y Social",
  "oat_num": "Solo el numero del OAT (ej: OA 6)",
  "oat_texto": "Texto completo del OAT elegido sin incluir el numero",
  "inicio": "Momento de inicio detallado (6-8 oraciones). La educadora presenta los materiales con entusiasmo usando dialogo directo con los ninos en primera persona (ej: 'Ninos y ninas, hoy les traje...'). Describe brevemente que es cada material de forma simple y cercana. Luego la educadora modela la accion ella primero ('Yo voy a probar... miren lo que pasa...'). Finaliza con una invitacion concreta a los ninos a participar ('Ahora ustedes pueden...'). Tono cercano, motivador y adecuado para 3 anos.",
  "escenario_1_titulo": "Titulo descriptivo escenario 1",
  "escenario_1_desc": "Descripcion detallada con mediacion y cierre 4-5 oraciones",
  "escenario_2_titulo": "Titulo descriptivo escenario 2",
  "escenario_2_desc": "Descripcion detallada con mediacion y cierre 4-5 oraciones",
  "escenario_3_titulo": "Titulo descriptivo escenario 3",
  "escenario_3_desc": "Descripcion detallada con mediacion y cierre 4-5 oraciones",
  "materiales_1": "material1, material2, material3, material4",
  "materiales_2": "material1, material2, material3",
  "materiales_3": "material1, material2, material3",
  "que_haran_ninos": "Descripcion en 4-5 oraciones en lenguaje simple y cercano para familias explicando que vivira el nino ese dia: que exploraran, con que materiales, que aprendizaje se busca. Tono calido y entusiasta, como si le contaras a un apoderado lo que hara su hijo.",
  "que_haran_familias": "Sugerencia de actividad familiar en casa en 4-5 oraciones: describe la actividad con pasos concretos, como invitar al nino a participar, que preguntas hacerle durante la actividad ('Que paso cuando...?', 'Como lo hiciste?') y como valorar lo que el nino hace o dice. Relacionada directamente con el tema del dia.",
  "que_necesitamos": "Lista de 3-5 materiales muy concretos y creativos que la familia puede conseguir facilmente en casa o en el comercio, relacionados directamente con el tema. Incluye una breve indicacion de como usarlos o prepararlos si es necesario. Nunca usar texto generico como 'materiales del hogar'.",
  "indicador_oa1": "Indicador observable 1 del OA comienza con verbo",
  "indicador_oa2": "Indicador observable 2 del OA comienza con verbo",
  "indicador_oat": "Indicador observable del OAT comienza con verbo"
}

IMPORTANTE: Usa solo comillas dobles. No uses saltos de linea dentro de los valores. Cada valor debe ser texto continuo en una sola linea."""


def build_prompt(data: dict) -> str:
    dias_texto = []
    for dia in data["dias"]:
        ideas_txt = ""
        if dia.get('ideas'):
            ideas_txt = "\nIdeas seleccionadas:\n" + "\n".join([
                f"- {idea['titulo']}{': ' + idea['descripcion'] if idea.get('descripcion') else ''}"
                for idea in dia['ideas']
            ])

        dias_texto.append(f"""{dia['nombre']}:
- Tema: {dia['tema']}{ideas_txt}""")

    return f"""Genera la planificacion semanal completa para Nivel Medio Menor.
Educadora: {data['educadora']}
Tecnicos: {data['tecnicos']}
Semana: {data['fecha_inicio']} al {data['fecha_fin']}

{chr(10).join(dias_texto)}

IMPORTANTE: Responde SOLO con JSON valido. Sin markdown. Sin saltos de linea dentro de los valores.
Elige el ambito, nucleo, OA y OAT mas adecuados para cada dia segun el tema e ideas."""


SUGGEST_PROMPT = """Eres una experta en Educación Parvularia chilena para Nivel Medio Menor (niños de 3 años).

Dado un tema del día, genera exactamente 5 ideas de experiencias de aprendizaje.
Responde SOLO con JSON válido sin markdown, sin texto adicional.
Formato:
{
  "ideas": [
    {"titulo": "Título corto de 4-6 palabras", "descripcion": "Descripción de 2-3 oraciones explicando la experiencia, materiales y aprendizaje esperado"},
    {"titulo": "...", "descripcion": "..."},
    {"titulo": "...", "descripcion": "..."},
    {"titulo": "...", "descripcion": "..."},
    {"titulo": "...", "descripcion": "..."}
  ]
}

IMPORTANTE: Títulos creativos y concretos. Descripciones en lenguaje pedagógico apropiado para párvulos. Sin saltos de línea dentro de los valores."""


# ═══════════════════════════════════════════════════════
# CONSTANTES
# ═══════════════════════════════════════════════════════

CORS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type, X-Groq-Key"
}

COLOR_HEADER = "ADDB7B"
COLOR_GRAY   = "F2F2F2"

DIAS_KEYS  = ["lunes", "martes", "miercoles", "jueves", "viernes"]
DIAS_NAMES = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]

ESCALA = ("Escala de apreciación:\n"
          "L: Logrado  ML: Medianamente Logrado\n"
          "IC: Iniciando Conducta  N/O: No observado  A: Ausente")

ROL_TEXTO = ("- Los niños/as escogen libremente su lugar de juego y material.\n"
             "- Los niños y niñas exploran y manipulan el material libremente.\n"
             "- Participan activamente durante la experiencia de aprendizaje.\n"
             "- Conversan y comparten sus ideas con sus pares y con los adultos.")

NINOS = [
    "RIHANNY ABOGADO","MAXIMILIANO AGÜERO","AMARO ALONSO","PAZ AMPUERO",
    "OLIVIA ANDRES","ANNY AULAR","CHRIS CARDENAS","ADIEL CARVAJAL",
    "GADIEL CARVAJAL","AMARU CASTRO","AURORA CERNA","MÍA CONTRERAS",
    "SOFÍA CORNEJO","ELOÍSA CORRALES","INTI FREDES","MÁXIMO GARCÍA",
    "DYLAN GELVIS","NICOLÁS GONZÁLEZ","TOMÁS GONZÁLEZ","EZEQUIEL HERRERA",
    "NOAH LARA","THIAGO LIZANO","BASTIÁN QUISPE","LUCIANO ROCCO",
    "MATEO RODRÍGUEZ","BRUNO SAAVEDRA","MIRANDA SAENZ","ARLETH SUÁREZ",
    "EMILIANO VARAS","ÓSCAR VILLASECA"
]

# Colores xlsx
C_L  = "C6EFCE"
C_ML = "FFEB9C"
C_IC = "FFCC99"
C_NO = "D9D9D9"
C_A  = "FCB4B4"


# ═══════════════════════════════════════════════════════
# UTILIDADES COMUNES
# ═══════════════════════════════════════════════════════

def clean_json(raw: str) -> dict:
    text = raw.strip().replace("```json", "").replace("```", "").strip()
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        text = match.group()
    result = []
    in_string = False
    escape = False
    for ch in text:
        if escape:
            result.append(ch); escape = False
        elif ch == '\\':
            result.append(ch); escape = True
        elif ch == '"':
            in_string = not in_string; result.append(ch)
        elif in_string:
            if ch in ('\n', '\r'): result.append(' ')
            elif ch == '\t':       result.append(' ')
            elif ord(ch) < 0x20:   pass
            else:                  result.append(ch)
        else:
            result.append(ch)
    return json.loads(''.join(result))


def call_groq(groq_key: str, body: dict) -> dict:
    client = Groq(api_key=groq_key)
    resp   = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        max_tokens=6000,
        temperature=0.7,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": build_prompt(body)}
        ]
    )
    return clean_json(resp.choices[0].message.content)


# ═══════════════════════════════════════════════════════
# UTILIDADES DOCX
# ═══════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    shading = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def cell_write(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return para


# ═══════════════════════════════════════════════════════
# UTILIDADES XLSX
# ═══════════════════════════════════════════════════════

def hfill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def hfont(bold=False, size=9, color="000000"):
    return Font(bold=bold, size=size, color=color, name="Arial")

def xcenter():
    return Alignment(horizontal="center", vertical="center", wrap_text=True)

def xleft():
    return Alignment(horizontal="left", vertical="center", wrap_text=True)

def thin_border():
    s = Side(style="thin", color="AAAAAA")
    return Border(left=s, right=s, top=s, bottom=s)


# ═══════════════════════════════════════════════════════
# BUILDER PLANIFICACIÓN DOCX
# ═══════════════════════════════════════════════════════

def build_planificacion(data: dict, ai: dict) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.orientation   = WD_ORIENT.LANDSCAPE
    section.page_width    = Cm(27.94)
    section.page_height   = Cm(21.59)
    section.top_margin    = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin   = Cm(1.27)
    section.right_margin  = Cm(1.27)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("PLANIFICACIÓN CO-CONSTRUIDA Y DIVERSIFICADA")
    r.bold = True; r.font.size = Pt(13)

    info1 = doc.add_paragraph()
    r = info1.add_run("\tNivel: "); r.bold = True; r.font.size = Pt(11)
    info1.add_run("Medio Menor").font.size = Pt(11)
    r2 = info1.add_run("                                                 Fecha: ")
    r2.bold = True; r2.font.size = Pt(11)
    info1.add_run(f"{data['fecha_inicio']} al {data['fecha_fin']}").font.size = Pt(11)

    info2 = doc.add_paragraph()
    r3 = info2.add_run("\tEducadora: "); r3.bold = True; r3.font.size = Pt(11)
    info2.add_run(data['educadora']).font.size = Pt(11)
    r4 = info2.add_run("                                 Técnicos: ")
    r4.bold = True; r4.font.size = Pt(11)
    info2.add_run(data['tecnicos']).font.size = Pt(11)

    doc.add_paragraph()

    C = [Cm(3.74), Cm(4.78), Cm(9.21), Cm(9.21), Cm(7.79), Cm(7.79), Cm(7.79)]
    table = doc.add_table(rows=13, cols=7)
    table.style = "Table Grid"
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = C[i]

    # Header OA
    r0 = table.rows[0]
    for i in range(7):
        set_cell_bg(r0.cells[i], COLOR_HEADER)
    cell_write(r0.cells[0], "ÁMBITO", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_write(r0.cells[1], "NÚCLEO", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    oa_hdr = r0.cells[2].merge(r0.cells[3])
    set_cell_bg(oa_hdr, COLOR_HEADER)
    cell_write(oa_hdr, "OBJETIVO DE APRENDIZAJE", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    oat_hdr = r0.cells[4].merge(r0.cells[5]).merge(r0.cells[6])
    set_cell_bg(oat_hdr, COLOR_HEADER)
    cell_write(oat_hdr, "OBJETIVO DE APRENDIZAJE TRANSVERSAL", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Filas OA por día
    for idx, key in enumerate(DIAS_KEYS):
        fd         = data["dias"][idx]
        aid        = ai.get(key, {})
        row        = table.rows[idx + 1]

        ambito_txt = aid.get("ambito", "")
        nucleo_txt = aid.get("nucleo", "")
        oa_txt     = f"{aid.get('oa_num','')}: {aid.get('oa_texto','')}"
        oat_nucleo = aid.get("oat_nucleo", "")
        oat_num    = aid.get("oat_num", "")
        oat_texto  = aid.get("oat_texto", "")

        c0 = row.cells[0]
        c0.paragraphs[0].clear()
        rb = c0.paragraphs[0].add_run(f"{DIAS_NAMES[idx]}: ")
        rb.bold = True; rb.font.size = Pt(9)
        c0.paragraphs[0].add_run(ambito_txt).font.size = Pt(9)

        cell_write(row.cells[1], nucleo_txt, size=9)

        oa_c = row.cells[2].merge(row.cells[3])
        cell_write(oa_c, oa_txt, size=9)

        oat_c = row.cells[4].merge(row.cells[5]).merge(row.cells[6])
        oat_c.paragraphs[0].clear()
        ra = oat_c.paragraphs[0].add_run("Ámbito: ")
        ra.bold = True; ra.font.size = Pt(9)
        oat_c.paragraphs[0].add_run("Desarrollo Personal y Social").font.size = Pt(9)
        pn = oat_c.add_paragraph()
        rn = pn.add_run("Núcleo: "); rn.bold = True; rn.font.size = Pt(9)
        pn.add_run(oat_nucleo).font.size = Pt(9)
        po = oat_c.add_paragraph()
        ro = po.add_run(f"{oat_num}: "); ro.bold = True; ro.font.size = Pt(9)
        po.add_run(oat_texto).font.size = Pt(9)

    # Header metodología
    r6 = table.rows[6]
    for i in range(7):
        set_cell_bg(r6.cells[i], COLOR_HEADER)
    cell_write(r6.cells[0], "ROL PROTAGÓNICO DEL NIÑO Y LA NIÑA", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    met_hdr = r6.cells[1].merge(r6.cells[2])
    set_cell_bg(met_hdr, COLOR_HEADER)
    cell_write(met_hdr, "SUGERENCIA METODOLÓGICA", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    mat_hdr = r6.cells[3].merge(r6.cells[4])
    set_cell_bg(mat_hdr, COLOR_HEADER)
    cell_write(mat_hdr, "RECURSOS/MATERIALES", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_write(r6.cells[5], "EVALUACIÓN", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_write(r6.cells[6], "PARTICIPACIÓN DE LA FAMILIA", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Filas metodología
    for idx, key in enumerate(DIAS_KEYS):
        fd  = data["dias"][idx]
        aid = ai.get(key, {})
        row = table.rows[idx + 7]

        c0 = row.cells[0]
        c0.paragraphs[0].clear()
        for i, line in enumerate(ROL_TEXTO.split("\n")):
            p = c0.paragraphs[0] if i == 0 else c0.add_paragraph()
            p.add_run(line).font.size = Pt(9)

        mc = row.cells[1].merge(row.cells[2])
        mc.paragraphs[0].clear()
        rt = mc.paragraphs[0].add_run(aid.get("titulo_dia", f"{DIAS_NAMES[idx]}: {fd.get('tema','')}"))
        rt.bold = True; rt.font.size = Pt(9)
        p_ini = mc.add_paragraph()
        ri = p_ini.add_run("Inicio: "); ri.bold = True; ri.font.size = Pt(9)
        p_ini.add_run(aid.get("inicio", "")).font.size = Pt(9)
        for s in ["1","2","3"]:
            pe = mc.add_paragraph()
            re_r = pe.add_run(f"Escenario {s}: {aid.get(f'escenario_{s}_titulo','')}")
            re_r.bold = True; re_r.font.size = Pt(9)
            pd = mc.add_paragraph()
            pd.add_run(aid.get(f"escenario_{s}_desc","")).font.size = Pt(9)

        matc = row.cells[3].merge(row.cells[4])
        matc.paragraphs[0].clear()
        first = True
        for s in ["1","2","3"]:
            p_sh = matc.paragraphs[0] if first else matc.add_paragraph()
            first = False
            rsh = p_sh.add_run(f"Escenario {s}:"); rsh.bold = True; rsh.font.size = Pt(9)
            for mat in [m.strip() for m in aid.get(f"materiales_{s}","").split(",") if m.strip()]:
                matc.add_paragraph().add_run(f"- {mat}").font.size = Pt(9)

        evc = row.cells[5]
        evc.paragraphs[0].clear()
        evc.paragraphs[0].add_run(ESCALA).font.size = Pt(8)
        ph = evc.add_paragraph()
        ph.add_run("Indicadores OA (Niños y Niñas):").bold = True
        ph.runs[0].font.size = Pt(9)
        for n, field in [("1","indicador_oa1"),("2","indicador_oa2")]:
            evc.add_paragraph().add_run(f"{n}. {aid.get(field,'')}").font.size = Pt(9)
        poh = evc.add_paragraph()
        poh.add_run("Indicador OAT:").bold = True
        poh.runs[0].font.size = Pt(9)
        evc.add_paragraph().add_run(f"1. {aid.get('indicador_oat','')}").font.size = Pt(9)

        famc = row.cells[6]
        famc.paragraphs[0].clear()
        famc.paragraphs[0].add_run(aid.get("que_haran_familias","")).font.size = Pt(9)

    # Fila observación
    r12 = table.rows[12]
    obs = r12.cells[0]
    for i in range(1, 7):
        obs = obs.merge(r12.cells[i])
    obs.paragraphs[0].clear()
    ro = obs.paragraphs[0].add_run("Observación:     ")
    ro.bold = True; ro.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════
# BUILDER INFO FAMILIA DOCX
# ═══════════════════════════════════════════════════════

def build_info_familia(data: dict, ai: dict) -> bytes:
    from datetime import datetime, timedelta

    COLOR_FAM_HEADER  = "185FA5"   # azul encabezado día
    COLOR_FAM_TEMA    = "EBF5FB"   # azul muy claro para fila tema
    COLOR_FAM_NINOS   = "E8F8F2"   # verde claro
    COLOR_FAM_FAM     = "FEF9E7"   # amarillo claro
    COLOR_FAM_MAT     = "F9EBEA"   # rojo claro

    doc = Document()
    section = doc.sections[0]
    section.orientation   = WD_ORIENT.LANDSCAPE
    section.page_width    = Cm(27.94)
    section.page_height   = Cm(21.59)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

    # Título principal
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title_p.add_run("¡Nos acompañas esta semana?")
    rt.bold = True; rt.font.size = Pt(14)
    rt.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)

    # Subtítulo con fechas
    try:
        fi = datetime.strptime(data["fecha_inicio"], "%Y-%m-%d")
        ff = datetime.strptime(data["fecha_fin"],    "%Y-%m-%d")
        MESES = ["enero","febrero","marzo","abril","mayo","junio",
                 "julio","agosto","septiembre","octubre","noviembre","diciembre"]
        rango = f"{fi.day} al {ff.day} de {MESES[ff.month-1]} de {ff.year}"
    except Exception:
        rango = f"{data.get('fecha_inicio','')} al {data.get('fecha_fin','')}"

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run(f"Semana del {rango}  ·  Nivel Medio Menor  ·  Educadora: {data.get('educadora','')}").font.size = Pt(10)

    doc.add_paragraph()

    # Tabla: fila encabezado + fila tema + 3 filas de contenido
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"

    try:
        fecha_ini = datetime.strptime(data["fecha_inicio"], "%Y-%m-%d")
        fechas = [(fecha_ini + timedelta(days=i)) for i in range(5)]
        dias_labels = [f"{DIAS_NAMES[i].upper()}\n{fechas[i].day}" for i in range(5)]
    except Exception:
        dias_labels = [n.upper() for n in DIAS_NAMES]

    # Fila 0: encabezados de día
    for i, label in enumerate(dias_labels):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, COLOR_FAM_HEADER)
        for line in label.split("\n"):
            p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            r.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Fila 1: tema del día
    for i, key in enumerate(DIAS_KEYS):
        aid = ai.get(key, {})
        tema_raw = data["dias"][i].get("tema", "") if i < len(data.get("dias",[])) else ""
        titulo   = aid.get("titulo_dia", tema_raw)
        cell = table.rows[1].cells[i]
        set_cell_bg(cell, COLOR_FAM_TEMA)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(titulo)
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x0C, 0x44, 0x7C)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas 2-4: contenido
    ROW_COLORS = [COLOR_FAM_NINOS, COLOR_FAM_FAM, COLOR_FAM_MAT]
    FIELDS = [
        ("que_haran_ninos",    "👧 ¿Qué harán los niños y niñas?"),
        ("que_haran_familias", "🏠 ¿Qué pueden hacer en casa?"),
        ("que_necesitamos",    "🛍 ¿Qué necesitamos?"),
    ]

    for i, key in enumerate(DIAS_KEYS):
        aid = ai.get(key, {})
        for row_idx, (field, label) in enumerate(FIELDS):
            cell = table.rows[row_idx + 2].cells[i]
            set_cell_bg(cell, ROW_COLORS[row_idx])
            cell.paragraphs[0].clear()
            # Label en negrita
            rl = cell.paragraphs[0].add_run(label + "\n")
            rl.bold = True; rl.font.size = Pt(8.5)
            # Contenido
            val = aid.get(field, "")
            cell.paragraphs[0].add_run(val).font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════
# BUILDER EVALUACIÓN XLSX
# ═══════════════════════════════════════════════════════

def build_evaluacion(form: dict, ai: dict) -> bytes:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    fecha_ini = form.get("fecha_inicio","")
    fecha_fin = form.get("fecha_fin","")
    educadora = form.get("educadora","Giovanna Marino")
    tecnicos  = form.get("tecnicos","Carol Miranda – Helen Oyarzún")

    # PORTADA
    ws_p = wb.create_sheet("Portada")
    ws_p.column_dimensions["A"].width = 35
    ws_p.column_dimensions["B"].width = 50

    ws_p.merge_cells("A1:B1")
    t = ws_p.cell(1,1,"EVALUACIÓN DE APRENDIZAJES")
    t.font = Font(bold=True, size=14, name="Arial", color="185FA5")
    t.alignment = xcenter()
    ws_p.row_dimensions[1].height = 28

    ws_p.merge_cells("A2:B2")
    ws_p.cell(2,1,f"Planificación Co-construida — Semana del {fecha_ini} al {fecha_fin}").font = hfont(size=11)
    ws_p.cell(2,1).alignment = xcenter()

    ws_p.merge_cells("A3:B3")
    ws_p.cell(3,1,"Nivel Medio Menor — Sala Cuna y Jardín Infantil Claudio Arrau").font = hfont(size=10)
    ws_p.cell(3,1).alignment = xcenter()

    ws_p.merge_cells("A5:B5")
    ws_p.cell(5,1,"IDENTIFICACIÓN").font = hfont(bold=True, size=10)
    ws_p.cell(5,1).fill = hfill(COLOR_HEADER)
    ws_p.cell(5,1).alignment = xcenter()

    for r, (lbl, val) in enumerate([
        ("Nivel:", "Medio Menor (2 a 3 años)"),
        ("Educadora de Párvulos:", educadora),
        ("Técnicas en Ed. Parvularia:", tecnicos),
        ("Período de evaluación:", f"{fecha_ini} al {fecha_fin}"),
        ("Total niños/as evaluados:", str(len(NINOS))),
        ("Instrumento:", "Escala de apreciación por indicador (OA y OAT)"),
    ], start=6):
        ws_p.cell(r,1,lbl).font = hfont(bold=True, size=10)
        ws_p.cell(r,1).alignment = xleft()
        ws_p.cell(r,2,val).font = hfont(size=10)
        ws_p.cell(r,2).alignment = xleft()

    ws_p.merge_cells("A13:B13")
    ws_p.cell(13,1,"ESCALA DE APRECIACIÓN").font = hfont(bold=True, size=10)
    ws_p.cell(13,1).fill = hfill(COLOR_HEADER)
    ws_p.cell(13,1).alignment = xcenter()

    for i, (cod, sig, desc, col) in enumerate([
        ("L",   "Logrado",               "El indicador se observa de forma autónoma y constante.", C_L),
        ("ML",  "Medianamente Logrado",  "Se observa con apoyo verbal o gestual del adulto.", C_ML),
        ("IC",  "Iniciando Conducta",    "Se observan acercamientos iniciales o esporádicos.", C_IC),
        ("N/O", "No Observado",          "No se logró observar durante la experiencia.", C_NO),
        ("A",   "Ausente",               "El niño/a no asistió a la jornada.", C_A),
    ]):
        r = 14 + i
        ws_p.cell(r,1,f"{cod} — {sig}").font = hfont(bold=True, size=9)
        ws_p.cell(r,1).fill = hfill(col)
        ws_p.cell(r,1).alignment = xleft()
        ws_p.cell(r,2,desc).font = hfont(size=9)
        ws_p.cell(r,2).alignment = xleft()

    # HOJAS DIARIAS
    day_sheets = {}
    for d_idx in range(5):
        key    = DIAS_KEYS[d_idx]
        fd     = form["dias"][d_idx]
        aid    = ai.get(key, {})
        tema   = aid.get("titulo_dia") or fd.get("tema","")
        ind1   = aid.get("indicador_oa1","")
        ind2   = aid.get("indicador_oa2","")
        ind_oat= aid.get("indicador_oat","")

        sname = f"{DIAS_NAMES[d_idx][:4]} {d_idx+1}"
        ws = wb.create_sheet(sname)
        day_sheets[d_idx] = sname

        ws.column_dimensions["A"].width = 5
        ws.column_dimensions["B"].width = 28
        ws.column_dimensions["C"].width = 32
        ws.column_dimensions["D"].width = 32
        ws.column_dimensions["E"].width = 32
        ws.column_dimensions["F"].width = 30

        ws.merge_cells("A1:F1")
        ws.cell(1,1,f"EVALUACIÓN — {DIAS_NAMES[d_idx]} {fecha_ini}: \"{tema}\"")
        ws.cell(1,1).font = Font(bold=True, size=12, name="Arial", color="185FA5")
        ws.cell(1,1).alignment = xcenter()
        ws.cell(1,1).fill = hfill("EFF6FF")
        ws.row_dimensions[1].height = 22

        meta = [
            ("Ámbito OA:",     aid.get("ambito","")),
            ("Núcleo OA:",     aid.get("nucleo","")),
            ("OA principal:",  f"{aid.get('oa_num','')}: {aid.get('oa_texto','')}"),
            ("Ámbito OAT:",    "Desarrollo Personal y Social"),
            ("Núcleo OAT:",    aid.get("oat_nucleo","")),
            ("OAT:",           f"{aid.get('oat_num','')}: {aid.get('oat_texto','')}"),
        ]
        for i, (lbl, val) in enumerate(meta):
            r = i + 2
            ws.merge_cells(f"A{r}:B{r}")
            ws.cell(r,1,lbl).font = hfont(bold=True, size=9)
            ws.cell(r,1).alignment = xleft()
            ws.merge_cells(f"C{r}:F{r}")
            ws.cell(r,3,val).font = hfont(size=9)
            ws.cell(r,3).alignment = xleft()

        for ci, (h, fc, cc) in enumerate(zip(
            ["N°","Niño/a","INDICADOR OA 1","INDICADOR OA 2","INDICADOR OAT","Observaciones"],
            [COLOR_HEADER,COLOR_HEADER,COLOR_HEADER,COLOR_HEADER,COLOR_HEADER,"185FA5"],
            ["000000","000000","000000","000000","000000","FFFFFF"]
        )):
            cell = ws.cell(8, ci+1, h)
            cell.font = hfont(bold=True, size=9, color=cc)
            cell.fill = hfill(fc)
            cell.alignment = xcenter()
            cell.border = thin_border()
        ws.row_dimensions[8].height = 18

        ws.cell(9,1,"").border = thin_border()
        ws.cell(9,2,"").border = thin_border()
        for ci, ind in enumerate([ind1, ind2, ind_oat]):
            lbl = f"OA-Ind {ci+1}: {ind}" if ci < 2 else f"OAT-Ind 1: {ind}"
            cell = ws.cell(9, ci+3, lbl)
            cell.font = hfont(size=8)
            cell.fill = hfill(COLOR_GRAY)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            cell.border = thin_border()
        ws.cell(9,6,"").fill = hfill(COLOR_GRAY)
        ws.cell(9,6).border = thin_border()
        ws.row_dimensions[9].height = 50

        dv = DataValidation(type="list", formula1='"L,ML,IC,N/O,A"', allow_blank=True)
        dv.error = "Ingrese: L, ML, IC, N/O o A"
        dv.errorTitle = "Valor inválido"
        dv.prompt = "Escala: L / ML / IC / N/O / A"
        ws.add_data_validation(dv)

        for n_idx, nombre_nino in enumerate(NINOS):
            r = n_idx + 10
            ws.cell(r,1,n_idx+1).font = hfont(size=9)
            ws.cell(r,1).alignment = xcenter()
            ws.cell(r,1).border = thin_border()
            ws.cell(r,2,nombre_nino).font = hfont(size=9)
            ws.cell(r,2).alignment = xleft()
            ws.cell(r,2).border = thin_border()
            ws.row_dimensions[r].height = 16
            for ci in range(3,6):
                cell = ws.cell(r, ci, "")
                cell.font = hfont(bold=True, size=10)
                cell.alignment = xcenter()
                cell.border = thin_border()
                dv.add(f"{get_column_letter(ci)}{r}")
                col_letter = get_column_letter(ci)
                for cod, color in [("L",C_L),("ML",C_ML),("IC",C_IC),("N/O",C_NO),("A",C_A)]:
                    ws.conditional_formatting.add(
                        f"{col_letter}{r}",
                        CellIsRule(operator="equal", formula=[f'"{cod}"'], fill=hfill(color))
                    )
            ws.cell(r,6,"").font = hfont(size=8)
            ws.cell(r,6).alignment = xleft()
            ws.cell(r,6).border = thin_border()

        for t_idx, (lbl, cod) in enumerate([
            ("TOTAL Logrado (L)","L"),
            ("TOTAL Medianamente Logrado (ML)","ML"),
            ("TOTAL Iniciando Conducta (IC)","IC"),
            ("TOTAL No Observado (N/O)","N/O"),
            ("TOTAL Ausentes (A)","A"),
        ]):
            r = 40 + t_idx
            ws.merge_cells(f"A{r}:B{r}")
            cell = ws.cell(r,1,lbl)
            cell.font = hfont(bold=True, size=9)
            cell.fill = hfill(COLOR_GRAY)
            cell.alignment = xleft()
            cell.border = thin_border()
            for ci in range(3,6):
                col = get_column_letter(ci)
                tc = ws.cell(r, ci, f'=COUNTIF({col}10:{col}39,"{cod}")')
                tc.font = hfont(bold=True, size=9)
                tc.alignment = xcenter()
                tc.border = thin_border()
                tc.fill = hfill(COLOR_GRAY)
            ws.cell(r,6).border = thin_border()

        ws.merge_cells("A45:B45")
        ws.cell(45,1,"% LOGRO (L+ML)").font = hfont(bold=True, size=9, color="FFFFFF")
        ws.cell(45,1).fill = hfill("185FA5")
        ws.cell(45,1).alignment = xleft()
        ws.cell(45,1).border = thin_border()
        for ci in range(3,6):
            col = get_column_letter(ci)
            formula = (f'=IFERROR((COUNTIF({col}10:{col}39,"L")+COUNTIF({col}10:{col}39,"ML"))'
                       f'/(COUNTA({col}10:{col}39)-COUNTIF({col}10:{col}39,"A"))*100,0)')
            tc = ws.cell(45, ci, formula)
            tc.number_format = '0.0"%"'
            tc.font = hfont(bold=True, size=9, color="FFFFFF")
            tc.fill = hfill("185FA5")
            tc.alignment = xcenter()
            tc.border = thin_border()
        ws.cell(45,6).border = thin_border()

    # RESUMEN SEMANAL
    ws_r = wb.create_sheet("Resumen Semanal")
    ws_r.column_dimensions["A"].width = 5
    ws_r.column_dimensions["B"].width = 28

    ws_r.merge_cells("A1:AC1")
    ws_r.cell(1,1,"RESUMEN SEMANAL POR NIÑO/A — Conteo de la escala")
    ws_r.cell(1,1).font = Font(bold=True, size=12, name="Arial", color="FFFFFF")
    ws_r.cell(1,1).fill = hfill("185FA5")
    ws_r.cell(1,1).alignment = xcenter()

    day_cols = [3, 8, 13, 18, 23]
    for d_idx in range(5):
        sc = day_cols[d_idx]
        ws_r.merge_cells(start_row=2, start_column=sc, end_row=2, end_column=sc+4)
        cell = ws_r.cell(2, sc, DIAS_NAMES[d_idx])
        cell.font = hfont(bold=True, size=9)
        cell.fill = hfill(COLOR_HEADER)
        cell.alignment = xcenter()
        cell.border = thin_border()
        for si, sub in enumerate(["L","ML","IC","N/O","A"]):
            c = ws_r.cell(3, sc+si, sub)
            c.font = hfont(bold=True, size=8)
            c.fill = hfill(COLOR_GRAY)
            c.alignment = xcenter()
            c.border = thin_border()

    for ci, hdr in [(1,"N°"),(2,"Niño/a"),(28,"Total Ind."),(29,"% Logro")]:
        ws_r.merge_cells(start_row=2, start_column=ci, end_row=3, end_column=ci)
        cell = ws_r.cell(2, ci, hdr)
        cell.font = hfont(bold=True, size=9, color="FFFFFF")
        cell.fill = hfill("185FA5")
        cell.alignment = xcenter()
        cell.border = thin_border()

    for n_idx, nombre_nino in enumerate(NINOS):
        r = n_idx + 4
        ws_r.cell(r,1,n_idx+1).font = hfont(size=9)
        ws_r.cell(r,1).alignment = xcenter()
        ws_r.cell(r,1).border = thin_border()
        ws_r.cell(r,2,nombre_nino).font = hfont(size=9)
        ws_r.cell(r,2).alignment = xleft()
        ws_r.cell(r,2).border = thin_border()

        for d_idx in range(5):
            sname = day_sheets[d_idx]
            sc    = day_cols[d_idx]
            nr    = n_idx + 10
            for si, cod in enumerate(["L","ML","IC","N/O","A"]):
                formula = f"=COUNTIF('{sname}'!C{nr}:E{nr},\"{cod}\")"
                cell = ws_r.cell(r, sc+si, formula)
                cell.font = hfont(size=9)
                cell.alignment = xcenter()
                cell.border = thin_border()

        ws_r.cell(r,28,f"=COUNTA(C{r}:AA{r})-COUNTIF(C{r}:AA{r},\"A\")").font = hfont(size=9)
        ws_r.cell(r,28).alignment = xcenter()
        ws_r.cell(r,28).border = thin_border()

        pct = f"=IFERROR((COUNTIF(C{r}:AA{r},\"L\")+COUNTIF(C{r}:AA{r},\"ML\"))/AB{r}*100,0)"
        ws_r.cell(r,29,pct).font = hfont(size=9)
        ws_r.cell(r,29).number_format = '0.0"%"'
        ws_r.cell(r,29).alignment = xcenter()
        ws_r.cell(r,29).border = thin_border()

    r_tot = len(NINOS) + 4
    ws_r.merge_cells(f"A{r_tot}:B{r_tot}")
    ws_r.cell(r_tot,1,"TOTAL NIVEL").font = hfont(bold=True, size=9, color="FFFFFF")
    ws_r.cell(r_tot,1).fill = hfill("185FA5")
    ws_r.cell(r_tot,1).alignment = xcenter()
    ws_r.cell(r_tot,1).border = thin_border()
    for ci in range(3, 30):
        col = get_column_letter(ci)
        cell = ws_r.cell(r_tot, ci, f"=SUM({col}4:{col}{r_tot-1})")
        cell.font = hfont(bold=True, size=9, color="FFFFFF")
        cell.fill = hfill("185FA5")
        cell.alignment = xcenter()
        cell.border = thin_border()

    # RESUMEN POR NÚCLEO
    ws_n = wb.create_sheet("Resumen por Núcleo")
    for ci, w in enumerate([8,35,30,30,8,8,8,8,8,10], 1):
        ws_n.column_dimensions[get_column_letter(ci)].width = w

    ws_n.merge_cells("A1:J1")
    ws_n.cell(1,1,"RESUMEN POR ÁMBITO Y NÚCLEO BCEP — Distribución de logros semanales")
    ws_n.cell(1,1).font = Font(bold=True, size=12, name="Arial", color="FFFFFF")
    ws_n.cell(1,1).fill = hfill("185FA5")
    ws_n.cell(1,1).alignment = xcenter()

    for ci, hdr in enumerate(["N°","Día y experiencia","Ámbito BCEP","Núcleo BCEP","L","ML","IC","N/O","A","% Logro"],1):
        cell = ws_n.cell(2, ci, hdr)
        cell.font = hfont(bold=True, size=9)
        cell.fill = hfill(COLOR_HEADER)
        cell.alignment = xcenter()
        cell.border = thin_border()

    for d_idx in range(5):
        key   = DIAS_KEYS[d_idx]
        fd    = form["dias"][d_idx]
        aid   = ai.get(key,{})
        tema  = aid.get("titulo_dia") or fd.get("tema","")
        sname = day_sheets[d_idx]
        r_oa  = 3 + d_idx*2

        for r, is_oat in [(r_oa, False),(r_oa+1, True)]:
            ws_n.cell(r,1,f"{d_idx+1}.{'OAT' if is_oat else 'OA'}").font = hfont(size=9)
            ws_n.cell(r,1).alignment = xcenter()
            ws_n.cell(r,1).border = thin_border()
            ws_n.cell(r,2,f"{DIAS_NAMES[d_idx]} — {tema}").font = hfont(size=9)
            ws_n.cell(r,2).alignment = xleft()
            ws_n.cell(r,2).border = thin_border()
            ws_n.cell(r,3,"Desarrollo Personal y Social (OAT)" if is_oat else aid.get("ambito","")).font = hfont(size=9)
            ws_n.cell(r,3).alignment = xleft()
            ws_n.cell(r,3).border = thin_border()
            ws_n.cell(r,4,aid.get("oat_nucleo","") if is_oat else aid.get("nucleo","")).font = hfont(size=9)
            ws_n.cell(r,4).alignment = xleft()
            ws_n.cell(r,4).border = thin_border()
            col_range = f"E10:E39" if is_oat else f"C10:D39"
            for si, cod in enumerate(["L","ML","IC","N/O","A"]):
                formula = f"=COUNTIF('{sname}'!{col_range},\"{cod}\")"
                cell = ws_n.cell(r, 5+si, formula)
                cell.font = hfont(size=9)
                cell.alignment = xcenter()
                cell.border = thin_border()
            pct = f"=IFERROR((E{r}+F{r})/(E{r}+F{r}+G{r}+H{r})*100,0)"
            ws_n.cell(r,10,pct).font = hfont(size=9)
            ws_n.cell(r,10).number_format = '0.0"%"'
            ws_n.cell(r,10).alignment = xcenter()
            ws_n.cell(r,10).border = thin_border()

    r_tot_n = 13
    ws_n.merge_cells(f"A{r_tot_n}:D{r_tot_n}")
    ws_n.cell(r_tot_n,1,"TOTAL").font = hfont(bold=True, size=9, color="FFFFFF")
    ws_n.cell(r_tot_n,1).fill = hfill("185FA5")
    ws_n.cell(r_tot_n,1).alignment = xcenter()
    ws_n.cell(r_tot_n,1).border = thin_border()
    for ci in range(5,11):
        col = get_column_letter(ci)
        cell = ws_n.cell(r_tot_n, ci, f"=SUM({col}3:{col}{r_tot_n-1})")
        cell.font = hfont(bold=True, size=9, color="FFFFFF")
        cell.fill = hfill("185FA5")
        cell.alignment = xcenter()
        cell.border = thin_border()

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════

@app.route(route="generate", methods=["GET","POST","OPTIONS"])
def generate(req: func.HttpRequest) -> func.HttpResponse:
    if req.method == "OPTIONS":
        return func.HttpResponse(status_code=200, headers=CORS)
    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Body JSON invalido", status_code=400, headers=CORS)
    groq_key = req.headers.get("X-Groq-Key") or os.environ.get("GROQ_API_KEY","")
    if not groq_key:
        return func.HttpResponse("Falta Groq API Key", status_code=401, headers=CORS)
    try:
        ai_data = call_groq(groq_key, body)
    except Exception as e:
        logging.error(f"Groq error: {e}")
        return func.HttpResponse(f"Error generando contenido: {e}", status_code=500, headers=CORS)
    out_hdrs = {**CORS, "Content-Type": "application/json"}
    return func.HttpResponse(
        body=json.dumps({"ok": True, "data": ai_data}, ensure_ascii=False),
        status_code=200, headers=out_hdrs
    )


@app.route(route="download-planificacion", methods=["POST","OPTIONS"])
def download_planificacion(req: func.HttpRequest) -> func.HttpResponse:
    if req.method == "OPTIONS":
        return func.HttpResponse(status_code=200, headers=CORS)
    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Body JSON invalido", status_code=400, headers=CORS)
    try:
        docx_bytes = build_planificacion(body["form"], body["ai"])
    except Exception as e:
        logging.error(f"DOCX planificacion error: {e}")
        return func.HttpResponse(f"Error generando planificacion: {e}", status_code=500, headers=CORS)
    fecha = body["form"].get("fecha_inicio","semana")
    out_hdrs = {**CORS, "Content-Disposition": f'attachment; filename="Planificacion_{fecha}.docx"'}
    return func.HttpResponse(body=docx_bytes, status_code=200,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=out_hdrs)


@app.route(route="download-familia", methods=["POST","OPTIONS"])
def download_familia(req: func.HttpRequest) -> func.HttpResponse:
    if req.method == "OPTIONS":
        return func.HttpResponse(status_code=200, headers=CORS)
    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Body JSON invalido", status_code=400, headers=CORS)
    try:
        docx_bytes = build_info_familia(body["form"], body["ai"])
    except Exception as e:
        logging.error(f"DOCX familia error: {e}")
        return func.HttpResponse(f"Error generando info familia: {e}", status_code=500, headers=CORS)
    fecha = body["form"].get("fecha_inicio","semana")
    out_hdrs = {**CORS, "Content-Disposition": f'attachment; filename="InfoFamilia_{fecha}.docx"'}
    return func.HttpResponse(body=docx_bytes, status_code=200,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=out_hdrs)


@app.route(route="download-evaluacion", methods=["POST","OPTIONS"])
def download_evaluacion(req: func.HttpRequest) -> func.HttpResponse:
    if req.method == "OPTIONS":
        return func.HttpResponse(status_code=200, headers=CORS)
    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Body JSON invalido", status_code=400, headers=CORS)
    try:
        xlsx_bytes = build_evaluacion(body["form"], body["ai"])
    except Exception as e:
        logging.error(f"XLSX evaluacion error: {e}")
        return func.HttpResponse(f"Error generando evaluacion: {e}", status_code=500, headers=CORS)
    fecha = body["form"].get("fecha_inicio","semana")
    out_hdrs = {**CORS, "Content-Disposition": f'attachment; filename="Evaluacion_{fecha}.xlsx"'}
    return func.HttpResponse(body=xlsx_bytes, status_code=200,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=out_hdrs)


@app.route(route="suggest-ideas", methods=["POST","OPTIONS"])
def suggest_ideas(req: func.HttpRequest) -> func.HttpResponse:
    if req.method == "OPTIONS":
        return func.HttpResponse(status_code=200, headers=CORS)
    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Body JSON invalido", status_code=400, headers=CORS)
    tema     = body.get("tema", "")
    dia      = body.get("dia", "")
    groq_key = os.environ.get("GROQ_API_KEY", "")
    if not groq_key:
        return func.HttpResponse("Falta GROQ_API_KEY", status_code=401, headers=CORS)
    try:
        client = Groq(api_key=groq_key)
        resp   = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=1000,
            temperature=0.8,
            messages=[
                {"role": "system", "content": SUGGEST_PROMPT},
                {"role": "user",   "content": f"Día: {dia}\nTema: {tema}\n\nGenera 5 ideas de experiencias de aprendizaje para niños de 3 años."}
            ]
        )
        ideas_data = clean_json(resp.choices[0].message.content)
    except Exception as e:
        logging.error(f"Groq suggest error: {e}")
        return func.HttpResponse(f"Error generando ideas: {e}", status_code=500, headers=CORS)
    out_hdrs = {**CORS, "Content-Type": "application/json"}
    return func.HttpResponse(
        body=json.dumps(ideas_data, ensure_ascii=False),
        status_code=200, headers=out_hdrs
    )