import logging
import json
import os
import io
from datetime import datetime

import azure.functions as func
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from groq import Groq

SYSTEM_PROMPT = """Eres una experta en Educación Parvularia chilena. Generas planificaciones semanales co-construidas y diversificadas para Nivel Medio Menor (3 años) siguiendo las Bases Curriculares de Educación Parvularia 2018.

Para cada día debes generar contenido pedagógico apropiado para niños de 3 años.

Responde SOLO con un objeto JSON válido sin texto adicional, sin bloques de código markdown, sin explicaciones previas.
El JSON debe tener exactamente estas claves por cada día (lunes, martes, miercoles, jueves, viernes):
{
  "lunes": {
    "inicio": "texto del momento de inicio (2-3 oraciones, actividad motivadora)",
    "escenario_1": "descripción del primer escenario de juego (3-4 oraciones detalladas)",
    "escenario_2": "descripción del segundo escenario de juego (3-4 oraciones detalladas)",
    "escenario_3": "descripción del tercer escenario de juego (3-4 oraciones detalladas)",
    "materiales": "material1, material2, material3, material4, material5",
    "participacion_familia": "sugerencia breve de participación familiar (1-2 oraciones)",
    "indicador_oa1": "primer indicador observable del OA (comienza con verbo, para niños de 3 años)",
    "indicador_oa2": "segundo indicador observable del OA (comienza con verbo, para niños de 3 años)",
    "indicador_oat": "indicador observable del OAT (comienza con verbo)"
  },
  "martes": { ... },
  "miercoles": { ... },
  "jueves": { ... },
  "viernes": { ... }
}"""


def build_prompt(data: dict) -> str:
    dias_texto = []
    for dia in data["dias"]:
        dias_texto.append(f"""{dia['nombre']}:
- Tema: {dia['tema']}
- Ámbito: {dia['ambito']}
- Núcleo: {dia['oa']['nucleo']}
- OA: {dia['oa']['num']} — {dia['oa']['texto']}
- OAT: {dia['oat']['num']} ({dia['oat']['nucleo']}) — {dia['oat']['texto']}""")

    return f"""Genera la planificación semanal completa con estos datos:

Educadora: {data['educadora']}
Técnicos: {data['tecnicos']}
Semana: {data['fecha_inicio']} al {data['fecha_fin']}

{chr(10).join(dias_texto)}

Genera contenido pedagógico rico, concreto y apropiado para niños de 3 años. Los escenarios deben ser distintos entre sí y promover el juego activo."""


def set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd {" ".join(f\'xmlns:{k}="{v}"\' for k, v in [("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")])} '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_header_row(table, texts: list, bg_color: str, text_color: str = "FFFFFF"):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_bg(cell, bg_color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(text_color)


def build_docx(data: dict, ai: dict) -> bytes:
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PLANIFICACIÓN CO-CONSTRUIDA Y DIVERSIFICADA")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)

    # Subtítulo
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Nivel: Medio Menor   |   Semana: {data['fecha_inicio']} al {data['fecha_fin']}").font.size = Pt(10)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.add_run(f"Educadora: {data['educadora']}   |   Técnicos: {data['tecnicos']}").font.size = Pt(10)

    doc.add_paragraph()

    dias_keys  = ["lunes", "martes", "miercoles", "jueves", "viernes"]
    dias_names = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]
    escala     = "L: Logrado   ML: Med. Logrado   IC: Iniciando Conducta   N/O: No observado   A: Ausente"

    for idx, key in enumerate(dias_keys):
        fd  = data["dias"][idx]
        aid = ai.get(key, {})

        # Título del día
        day_title = doc.add_paragraph()
        day_run = day_title.add_run(f"{dias_names[idx].upper()}: {fd['tema']}")
        day_run.bold = True
        day_run.font.size = Pt(11)
        day_run.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)

        # ── TABLA 1: OA ──
        t1 = doc.add_table(rows=2, cols=4)
        t1.style = "Table Grid"
        t1.autofit = False

        # Anchos columnas tabla 1
        widths1 = [Cm(3.5), Cm(4), Cm(7), Cm(5.5)]
        for row in t1.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths1[i]

        # Header tabla 1
        headers1 = ["ÁMBITO", "NÚCLEO", "OBJETIVO DE APRENDIZAJE", "OA TRANSVERSAL"]
        add_header_row(t1, headers1, "185FA5")

        # Datos tabla 1
        row1 = t1.rows[1]
        row1.cells[0].paragraphs[0].add_run(fd["ambito"]).font.size = Pt(9)
        row1.cells[1].paragraphs[0].add_run(fd["oa"]["nucleo"]).font.size = Pt(9)
        row1.cells[2].paragraphs[0].add_run(f"{fd['oa']['num']}: {fd['oa']['texto']}").font.size = Pt(9)
        row1.cells[3].paragraphs[0].add_run(f"{fd['oat']['num']}: {fd['oat']['texto']}").font.size = Pt(9)

        doc.add_paragraph()

        # ── TABLA 2: Metodología ──
        t2 = doc.add_table(rows=2, cols=5)
        t2.style = "Table Grid"
        t2.autofit = False

        widths2 = [Cm(2.5), Cm(7.5), Cm(3.5), Cm(4), Cm(2.5)]
        for row in t2.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths2[i]

        headers2 = ["ROL\nPROTAGÓNICO", "SUGERENCIA METODOLÓGICA", "RECURSOS Y\nMATERIALES", "EVALUACIÓN", "FAMILIA"]
        add_header_row(t2, headers2, "1D9E75")

        # Datos tabla 2
        row2 = t2.rows[1]

        # Rol protagónico
        row2.cells[0].paragraphs[0].add_run("Niño/a").font.size = Pt(9)

        # Sugerencia metodológica
        met_cell = row2.cells[1]
        met_cell.paragraphs[0].clear()
        p = met_cell.paragraphs[0]
        r = p.add_run(f"{dias_names[idx]} {data['fecha_inicio']}: '{fd['tema']}'")
        r.bold = True
        r.font.size = Pt(9)

        def add_scenario(cell, label, text):
            p2 = cell.add_paragraph()
            r2 = p2.add_run(f"{label}: ")
            r2.bold = True
            r2.font.size = Pt(9)
            r3 = p2.add_run(text)
            r3.font.size = Pt(9)

        add_scenario(met_cell, "Inicio", aid.get("inicio", ""))
        add_scenario(met_cell, "E1", aid.get("escenario_1", ""))
        add_scenario(met_cell, "E2", aid.get("escenario_2", ""))
        add_scenario(met_cell, "E3", aid.get("escenario_3", ""))

        # Materiales
        mat_cell = row2.cells[2]
        mat_cell.paragraphs[0].clear()
        for mat in aid.get("materiales", "").split(","):
            p3 = mat_cell.add_paragraph(mat.strip())
            p3.style.font.size = Pt(9) if p3.style else None
            for run in p3.runs:
                run.font.size = Pt(9)

        # Evaluación
        ev_cell = row2.cells[3]
        ev_cell.paragraphs[0].clear()
        ep = ev_cell.paragraphs[0]
        er = ep.add_run(escala)
        er.font.size = Pt(8)
        er.italic = True

        def add_indicator(cell, label, text):
            pi = cell.add_paragraph()
            ri1 = pi.add_run(f"{label}: ")
            ri1.bold = True
            ri1.font.size = Pt(9)
            ri2 = pi.add_run(text)
            ri2.font.size = Pt(9)

        add_indicator(ev_cell, "Ind. OA 1", aid.get("indicador_oa1", ""))
        add_indicator(ev_cell, "Ind. OA 2", aid.get("indicador_oa2", ""))
        add_indicator(ev_cell, "Ind. OAT",  aid.get("indicador_oat", ""))

        # Familia
        fam_cell = row2.cells[4]
        fam_cell.paragraphs[0].clear()
        fp2 = fam_cell.paragraphs[0]
        fp2.add_run(aid.get("participacion_familia", "")).font.size = Pt(9)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def main(req: func.HttpRequest) -> func.HttpResponse:
    # CORS preflight
    if req.method == "OPTIONS":
        return func.HttpResponse(
            status_code=200,
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type, X-Groq-Key"
            }
        )

    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Body JSON inválido", status_code=400)

    groq_key = req.headers.get("X-Groq-Key") or os.environ.get("GROQ_API_KEY", "")
    if not groq_key:
        return func.HttpResponse("Falta Groq API Key", status_code=401)

    try:
        client = Groq(api_key=groq_key)
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            max_tokens=4000,
            temperature=0.7,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user",   "content": build_prompt(body)}
            ]
        )
        content = response.choices[0].message.content.strip()
        content = content.replace("```json", "").replace("```", "").strip()
        ai_data = json.loads(content)
    except json.JSONDecodeError as e:
        logging.error(f"JSON parse error: {e} — content: {content}")
        return func.HttpResponse(f"Error parseando respuesta IA: {e}", status_code=500)
    except Exception as e:
        logging.error(f"Groq error: {e}")
        return func.HttpResponse(f"Error llamando a Groq: {e}", status_code=500)

    try:
        docx_bytes = build_docx(body, ai_data)
    except Exception as e:
        logging.error(f"DOCX error: {e}")
        return func.HttpResponse(f"Error generando .docx: {e}", status_code=500)

    fecha = body.get("fecha_inicio", "semana")
    filename = f"Planificacion_{fecha}.docx"

    return func.HttpResponse(
        body=docx_bytes,
        status_code=200,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Access-Control-Allow-Origin": "*"
        }
    )